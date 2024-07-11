#! python

"""Populates a leave form from the command line"""

# Python imports
import argparse
import datetime
from importlib import resources
import json
import os
import platform
import sys
from typing import Any, Optional

# Module imports
import appdirs
from docx import Document


def _get_response(k: str, v: Any) -> Any:
    happy_with_response = False
    while not happy_with_response:
        user_input = input(f"Enter {k} (press Enter to use {v}): ")
        user_input = user_input if user_input else v
        response = input(
            f"Is {user_input} correct? "
            "y to continue, n to re-enter or q to quit: "
        )
        happy_with_response = (
            response.lower() == 'y'
            or response.lower() == 'yes'
            or response.lower() == 'ydw'
        )
        if response.lower() == 'q' or response.lower() == 'quit':
            sys.exit()
    return user_input

def _format_date(iso_date: str) -> str:
    return "/".join(reversed(iso_date.split("-")))


def init(
        conf_path: Optional[str] = None,
        force_init: bool = False,
        update_balance: Optional[int] = None,
) -> dict:
    """Initilise config file

    If the file already exists, simply returns the config
    file as a dictionary.

    Args:
        conf_path (str, optional) : Path to config path. If None,

        force_init (bool, optional) : If True, will re-initialise
                the config file even if it already exists.
                Defaults to False.
        update_balance (bool, optional) : If not None, will update
                the balance to the value of update_balance.
                Defaults to balance.

    Returns:
        config (dict) : Configuration dictionary.
    """

    if conf_path is None:
        conf_dir = appdirs.user_config_dir('popleave')
        conf_path = os.path.join(conf_dir, "config.json")

        if not os.path.exists(conf_dir):
            os.makedirs(conf_dir)

    if not os.path.exists(conf_path) or force_init:

        config = {
            "name": "Samwise Gamgee",
            "department": "The Shire",
            "remaing_days_leave": 26,
        }

        for k, v in config.items():
            config[k] = _get_response(k, v)

        with open(conf_path, "w+", encoding = "utf-8") as fp:
            json.dump(config, fp)
            
    else:
        with open(conf_path, "r", encoding = "utf-8") as fp:
            config = json.load(fp)

        if update_balance is not None:
            with open(conf_path, "w", encoding = "utf-8") as fp:
                config["remaing_days_leave"] = update_balance
                json.dump(config, fp)

    return config


def get_doc_dict(
        config: dict,
        start_date: Optional[str | datetime.date] = None,
        duration: Optional[int | float | datetime.timedelta] = None,
        end_date: Optional[str | datetime.date] = None,
        reason: Optional[str] = None,
        is_toil: bool = False,
) -> dict:
    """Return populated document dictionary

    Args:
        config (dict): Config dictionary.
        start_date (str or datetime.date, optional): Leave start date.
                If None, will prompt user for a start_date.
                Defaults to None.
        duration (int or float or datetime.timedelta, optional) : Duration
                of leave. If None, will prompt user or a duration.
        end_date (str or datetime.date, optional) : Leave end date.
                If None, will prompt the user for an end_date.
                Defaults to None.
        reason (str, optional) : Reason for leave. If None, will prompt
                the user for a reason for leave.
        is_toil (bool, optional) : If True, fills the form out as a
                Time Off In Leu instead of Annual leave.
                Defaults to False.

    Returns:
        doc_dict (dict) : Dictionary of documentation fields.
    """

    doc_dict = dict()

    if start_date is None:
        start_date = _get_response(
            'Start date (YYYY-MM-DD)', str(datetime.date.today()),
        )
    start_date = datetime.date.fromisoformat(start_date)
    
    if duration is None and end_date is None:
        duration = int(_get_response('number of days', 1))
    
    if duration is None:
        duration = datetime.data.fromisoformat(end_date) - start_date
    else:
        duration = datetime.timedelta(days=duration)

    if end_date is None:
        end_date = start_date + duration

    if reason is None:
        reason = _get_response('reason for leave', 'Holiday')

    # Of the format (value, (table, row, col)))
    doc_dict['name'] = (config['name'], (0, 0, 0))
    doc_dict['department'] = (config['department'], (0, 0, 1))
    doc_dict['today'] = (
        _format_date(str(datetime.date.today())),
        (0, 0, 2),
    )
    
    type_row = 3 if is_toil else 1
    doc_dict['type'] = ('x', (1, type_row, 1))
    
    doc_dict['start_date'] = (
        _format_date(str(start_date)),
        (2, 3, 0),
    )
    doc_dict['end_date'] = (
        _format_date(str(end_date)),
        (2, 3, 1),
    )
    doc_dict['duration'] = (duration.days, (2, 3, 2))
    if is_toil:
        balance = config['remaing_days_leave']
    else:
        balance = int(config['remaing_days_leave']) - duration.days
    doc_dict['balance'] = (balance, (2, 3, 3))
        
    doc_dict['reason'] = (str(reason), (3, 0, 0))

    return doc_dict


def populate_file(
        doc_dict: dict,
        template_file: Optional[str] = None,
) -> bool:
    """Populate a word template file with annual leave info

    Args:
        doc_dict (dict) : Dictionary of leave info.
        template_file (str, optional) : Template file to populate.
                If None, will load the template.docx file associated with
                the package. Defaults to None.
    
    Returns:
        return_val (bool) : True if successful, False otherwise.
    """

    my_package = (
        os.path.splitext(os.path.basename(__file__))[0]
        if not __package__
        else __package__
    )
    template_file = (
        resources.files(my_package).joinpath('template.docx')
        if template_file is None
        else template_file
    )
    
    doc = Document(template_file)
    for value, (t, r, c) in doc_dict.values():
        doc.tables[t].cell(r, c).text = (
            f"{doc.tables[t].cell(r, c).text}\n{value}"
        )

    initials = "".join([n[0].upper() for n in doc_dict['name'][0].split()])
    startday = "".join(doc_dict['start_date'][0].split("/"))

    
    new_file_name = f"{initials}_ANNUAL_{startday}.docx"
    doc.save(new_file_name)

    new_config = init(update_balance=int(doc_dict["balance"][0]))

    print(f"You have {new_config['remaing_days_leave']} days holiday remaining.")
    
    return True


def main(
        force_init: bool = False, is_toil: bool = False, **kwargs: Any,
) -> bool:
    """Main script"""
    
    config = init(force_init=force_init)
    doc_dict = get_doc_dict(config, is_toil=is_toil, **kwargs)
    return populate_file(doc_dict)


if __name__ == "__main__":

    #####################
    # Formats Arguments #
    #####################
    
    parser = argparse.ArgumentParser(
        description='Populates the leave of absence form with a CLI',
        allow_abbrev=True,
    )

    parser.add_argument(
        '-i',
        '--init',
        action='store_true',
        help='Forces reinitialisation of config file.',
    )

    parser.add_argument(
        '-s',
        '--start_date',
        type=str,
        default=None,
        help='Start date in YYYY-MM-DD format.'
    )

    parser.add_argument(
        '-e',
        '--end_date',
        type=str,
        default=None,
        help='End date in YYYY-MM-DD format.'
    )

    parser.add_argument(
        '-d',
        '--duration',
        type=int,
        default=None,
        help='Duration in days.',
    )

    parser.add_argument(
        '-r',
        '--reason',
        type=str,
        default=None,
        help='Reason for leave.',
    )

    parser.add_argument(
        '-t',
        '--toil',
        action='store_true',
        help='Fills out Time Off In Leu instead of annual leave.',
    )

    args = vars(parser.parse_args())
    force_init = args.pop('init')
    is_toil = args.pop('toil')

    #####################
    # Runs main program #
    #####################

    main(force_init=force_init, is_toil=is_toil, **args)
